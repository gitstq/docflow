"""Word document to Markdown converter."""

from pathlib import Path

from docflow.core.converters.base import BaseConverter
from docflow.core.models import DocumentMetadata

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class WordConverter(BaseConverter):
    """Convert Word documents to Markdown."""

    def convert(self, source_path: Path, output_path: Path):
        """Convert a Word document to Markdown."""
        if not HAS_DOCX:
            return self._create_error_result(
                source_path,
                output_path,
                ["python-docx not installed. Run: pip install python-docx"],
            )

        try:
            doc = Document(source_path)
            markdown_parts = []
            images = []

            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect heading level
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        level_num = int(level)
                        text = "#" * min(level_num, 6) + " " + text
                    except ValueError:
                        pass
                elif para.style.name == "Title":
                    text = "# " + text

                # Handle bold/italic
                if para.runs:
                    formatted_parts = []
                    for run in para.runs:
                        run_text = run.text
                        if run.bold and run.italic:
                            run_text = f"***{run_text}***"
                        elif run.bold:
                            run_text = f"**{run_text}**"
                        elif run.italic:
                            run_text = f"*{run_text}*"
                        formatted_parts.append(run_text)
                    text = "".join(formatted_parts)

                markdown_parts.append(text)

            # Extract tables
            for table in doc.tables:
                md_table = self._format_table(table)
                markdown_parts.append(f"\n\n{md_table}\n")

            markdown_content = "\n\n".join(markdown_parts)
            metadata = self._extract_metadata(source_path, doc)

            return self._create_success_result(
                source_path,
                output_path,
                markdown_content,
                metadata=metadata,
                images=images,
            )

        except Exception as e:
            return self._create_error_result(
                source_path, output_path, [f"Word conversion error: {str(e)}"]
            )

    def _format_table(self, table) -> str:
        """Convert a Word table to Markdown."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        header = rows[0]
        separator = ["---"] * len(header)
        body = rows[1:]

        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(separator) + " |",
        ]
        for row in body:
            lines.append("| " + " | ".join(row) + " |")

        return "\n".join(lines)

    def _extract_metadata(self, source_path: Path, doc) -> DocumentMetadata:
        """Extract metadata from Word document."""
        props = doc.core_properties

        return DocumentMetadata(
            title=props.title,
            author=props.author,
            subject=props.subject,
            keywords=props.keywords.split(", ") if props.keywords else [],
            creator=props.category,
            creation_date=props.created,
            modification_date=props.modified,
            file_size=source_path.stat().st_size if source_path.exists() else None,
        )
